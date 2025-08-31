from __future__ import annotations
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
import json
import csv

def _ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)

def export_json(logs: List[Dict[str, Any]], filename: str) -> None:
    path = Path(filename)
    _ensure_parent(path)
    with path.open("w", encoding="utf-8") as f:
        json.dump(logs, f, ensure_ascii=False, indent=2, default=str)

def export_csv(logs: List[Dict[str, Any]], filename: str) -> None:
    if not logs:
        raise ValueError("No logs to export!")
    path = Path(filename)
    _ensure_parent(path)
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(logs[0].keys()))
        writer.writeheader()
        writer.writerows(logs)

def export_pdf(logs: List[Dict[str, Any]], filename: str) -> None:
    path = Path(filename)
    _ensure_parent(path)
    cnvs = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    y = height - 40
    for i, log in enumerate(logs, start=1):
        line = f"{log.get('id','-')} | {log.get('log_type','-')} | {log.get('log_message','')} | {log.get('hostname','-')} | {log.get('created_at','-')}"
        cnvs.drawString(30, y, line[:120])
        y -= 18
        if y < 40:
            cnvs.showPage()
            y = height - 40
    cnvs.save()

def export_docx(logs: List[Dict[str, Any]], filename: str) -> None:
    path = Path(filename)
    _ensure_parent(path)
    doc = Document()
    doc.add_heading("Log Export", level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "Type"
    hdr[2].text = "Message"
    hdr[3].text = "Hostname"
    hdr[4].text = "Created At"
    for log in logs:
        row = table.add_row().cells
        row[0].text = str(log.get("id", ""))
        row[1].text = str(log.get("log_type", ""))
        row[2].text = str(log.get("log_message", ""))
        row[3].text = str(log.get("hostname", ""))
        row[4].text = str(log.get("created_at", ""))
    doc.save(str(path))